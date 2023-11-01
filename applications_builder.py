import pandas as pd
from docx import Document


class ApplicationsBuilder():

    def __init__(self):
        """
        Data is the pd.DataFrame or pd.Series that
        contains all data for one user
        """
        self.data = pd.DataFrame()

    def _read_blank(self, path) -> str:
        """
        Reads the blank and transforms it into a text
        :param path:
        :return: str
        """
        document = Document(path)

        text = []
        for paragraph in document.paragraphs:
            text.append(paragraph.text)
        text = "\n".join(text)

        return text

    def _build_card_obtaining(self) -> Document():
        """
        Builds card obtaining application
        :return: Document()
        TODO: finish card obtaining
        """
        path = "card_obtaining_blank.docx"  # TODO: make blank
        text = self._read_blank(path)
        info = ...

    def _build_personal_data_processing_consent(self) -> Document():
        """
        Builds card obtaining application
        :return: Document()
        TODO: finish personal data processing consent
        """
        path = "personal_data_processing_consent.docx"  # TODO: make blank
        text = self._read_blank(path)
        info = ...

    def feed_data(self, data) -> None:
        """
        Feeds the data to the ApplicationsBuilder()
        :param data: pd.DataFrame(), pd.Series()
        :return: None
        """
        self.data = data

    def build_applications(self) -> None:
        """
        Builds all the applications using
        previously defined functions
        :return: None
        """
        name, surname, patronymic = \
            self.data.loc["Name"], self.data.loc["Surname"], self.data.loc["Patronymic"]

        card = self._build_card_obtaining()
        consent = self._build_personal_data_processing_consent()

        card.save(f"Applications/Заява на отримання Транспортної картки {name} {surname} {patronymic}")
        consent.save(f"Applications/Згода на обробку особистих даних {name} {surname} {patronymic}")


if __name__ == '__main__':
    builder = ApplicationsBuilder()
    data = ...
    builder.feed_data(data)
    builder.build_applications()
    builder.move_to_directory()
