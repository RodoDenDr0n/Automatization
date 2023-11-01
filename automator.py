import pandas as pd
from docx import Document


def process_picture():
    pass


def get_text(document):
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text)
    text = "\n".join(text)
    return text


def paste_data(text, data):
    for i in data.index:
        info = tuple(data.loc[i])
        parsed = text % info
        document = Document()
        for paragraph in parsed.split("\n"):
            document.add_paragraph(paragraph)
        document.save(f"Заяви/Заява_{info[0]}_{info[1]}.docx")


if __name__ == '__main__':
    data = pd.read_csv("Тестування отримання інформації (Responses) - Form Responses 1.csv").drop("Timestamp", axis=1)
    document = Document("test.docx")
    text = get_text(document)
    paste_data(text, data)
